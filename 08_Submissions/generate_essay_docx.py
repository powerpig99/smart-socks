"""Generate individual essay as a .docx with professional minimalistic formatting."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# --- Define styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading 1 style
h1_style = doc.styles["Heading 1"]
h1_style.font.name = "Calibri"
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
h1_style.paragraph_format.space_before = Pt(0)
h1_style.paragraph_format.space_after = Pt(4)

# Heading 2 style
h2_style = doc.styles["Heading 2"]
h2_style.font.name = "Calibri"
h2_style.font.size = Pt(12)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
h2_style.paragraph_format.space_before = Pt(14)
h2_style.paragraph_format.space_after = Pt(4)

# --- Title ---
title = doc.add_heading("Individual Essay", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- Subtitle / metadata line ---
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
meta.paragraph_format.space_after = Pt(2)
run = meta.add_run("ELEC-E7840 Smart Wearables  |  Spring 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.LEFT
meta2.paragraph_format.space_after = Pt(2)
run2 = meta2.add_run("Jing Liang  |  Smart Socks for Physical Activity Recognition (Topic 3)")
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Thin rule
rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(6)
rule.paragraph_format.space_after = Pt(6)
run_rule = rule.add_run("_" * 72)
run_rule.font.size = Pt(6)
run_rule.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# --- Helper to add a body paragraph with inline code formatting ---
def add_body(text):
    """Add a paragraph, rendering `code` segments in monospace."""
    p = doc.add_paragraph()
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:  # inside backticks
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


# --- Content ---

doc.add_heading("What I Learned", level=2)
add_body(
    "In the Smart Socks project, I handled most of the software and electronics: "
    "designing circuits, programming the ESP32-S3 firmware, building Python data "
    "collection and visualization tools, and implementing the communication stack "
    "(serial, WiFi, and BLE). This experience taught me that wearable system "
    "development is fundamentally different from pure software work \u2014 the hardware "
    "imposes constraints that remain invisible during desk testing and only reveal "
    "themselves when you test in the actual operating environment."
)

doc.add_heading("Challenge 1: ESP32-S3 Native USB Behavior", level=2)
add_body(
    "The most frustrating technical challenge was making the firmware work reliably "
    "across different conditions. The ESP32-S3 XIAO uses native USB rather than a "
    "separate UART bridge chip, which introduces behaviors that are poorly documented "
    "and difficult to diagnose."
)
add_body(
    "During development, everything worked fine on my laptop. But on battery power, "
    "the device froze at boot. The root cause was `while (!Serial) delay(10);` \u2014 a "
    "common Arduino pattern that waits for the serial port to be ready. On native USB, "
    "this condition is never satisfied without a connected computer, so the device "
    "blocked forever. The fix was trivial, but finding it required understanding the "
    "hardware-level difference between a UART bridge and the ESP32-S3\u2019s built-in USB CDC."
)
add_body(
    "A second issue appeared when the device was connected to USB but no terminal was "
    "open: `Serial.println()` at 50 Hz fills the transmit buffer in seconds, and once "
    "full, every write blocks the entire main loop \u2014 sensors stop, WiFi drops, "
    "everything halts. The solution was guarding all serial output with `if (Serial)` "
    "to skip writes when no host is listening."
)
add_body(
    "These bugs taught me to always test embedded devices in every deployment scenario "
    "(battery-only, USB with terminal open, USB without terminal), not just the "
    "convenient one at the workbench."
)

doc.add_heading("Challenge 2: WiFi That Blocks Everything", level=2)
add_body(
    "WiFi connectivity went through several iterations. The initial reconnection logic "
    "used `WiFi.scanNetworks()`, which blocks the processor for 2\u20135 seconds. At Aalto, "
    "with dozens of visible access points, the scan would find many networks but none "
    "matching our saved hotspot, causing repeated multi-second freezes in the sensor "
    "data stream."
)
add_body(
    "The fix was a tiered reconnection strategy: try `WiFi.reconnect()` first (nearly "
    "instant, reconnects to the last known AP), and only fall back to a full scan after "
    "repeated failures. More importantly, I adopted the principle that WiFi must never "
    "block core functionality \u2014 sensor sampling and serial streaming continue regardless "
    "of WiFi state. This single architectural decision prevented a whole class of "
    "reliability problems."
)
add_body(
    "I also discovered that Aalto\u2019s open WiFi blocks device-to-device HTTP, making it "
    "unusable for our web dashboard demo. We switched to a phone hotspot, which turned "
    "out to be more reliable and portable."
)

doc.add_heading("Challenge 3: Iterative Sensor Design", level=2)
add_body(
    "Our sensor configuration evolved substantially through the project. We explored "
    "different numbers and placements of sensors before arriving at the final design: "
    "3 pressure sensors per foot (heel, inner ball, outer ball) and 1 stretch sensor "
    "per knee, totaling 8 sensors across two ESP32 boards (one per leg)."
)
add_body(
    "The three-point pressure layout on the foot captures weight distribution more "
    "completely than two sensors could \u2014 the medial and lateral metatarsal separation "
    "helps distinguish activities like leaning or turning, where weight shifts sideways. "
    "The knee stretch sensor detects flexion angle, which is critical for separating "
    "stair climbing from flat walking. Each iteration taught us something about what "
    "information the system would actually need to distinguish our target activities."
)
add_body(
    "Building supporting tools was equally important. The real-time calibration "
    "visualizer I wrote \u2014 with live plots, GIF recording, and serial command "
    "integration \u2014 gave the whole team immediate visual feedback on whether sensors "
    "were responding correctly. This made integration sessions far more productive than "
    "interpreting raw numbers on a terminal."
)

doc.add_heading("Challenge 4: Team Dynamics", level=2)
add_body(
    "With three team members having different skills and availability, I ended up "
    "taking on the bulk of the electronics and software work while my teammates focused "
    "on sensor fabrication, garment design, and documentation. The main challenge was "
    "that hardware and firmware development are tightly coupled \u2014 I needed working "
    "sensors to test firmware, but sensor fabrication had its own timeline. I addressed "
    "this by using early prototype sensors I had made during a class session to develop "
    "and debug the firmware, then validating with the final textile sensors once they "
    "were ready. This decoupled the two workflows and avoided a situation where hardware "
    "and software problems would surface simultaneously."
)
add_body(
    "I also explored the machine learning side briefly, setting up the preprocessing "
    "and feature extraction pipeline. Even without training a model yet, this early "
    "exploration shaped practical decisions \u2014 choosing the right sampling rate, CSV "
    "format, and labeling conventions from the start so that data collected now would "
    "be directly usable for classification later."
)

doc.add_heading("Reflection", level=2)
add_body(
    "The most valuable lesson from this project is that wearable systems fail at "
    "boundaries: between hardware and software, between bench testing and real-world "
    "use, between teammates working on different subsystems. Technical skills like "
    "ESP32 programming, protocol implementation, and Python tooling are important, but "
    "the deeper insight is about process \u2014 always test in the real deployment "
    "environment, simplify the design wherever possible, and build tools that make the "
    "system\u2019s behavior visible to everyone. These principles will stay with me well "
    "beyond this course."
)

# --- Save ---
output_path = "/Users/jingliang/Documents/active_projects/Smart Socks/08_Submissions/individual_essay_jing.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
